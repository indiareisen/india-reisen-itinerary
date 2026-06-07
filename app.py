import streamlit as st
import json
from datetime import datetime, timedelta
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import folium
from streamlit_folium import st_folium

# Page config
st.set_page_config(
    page_title="India Reisen - Full Service Itinerary Generator",
    page_icon="🌸",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    :root {
        --gold: #D4A574;
        --magenta: #d1356f;
    }
    
    .main-header {
        color: var(--magenta);
        font-size: 2.5rem;
        font-weight: bold;
        margin-bottom: 0.5rem;
    }
    
    .sub-header {
        color: var(--gold);
        font-size: 1.2rem;
        margin-bottom: 2rem;
    }
    
    .card {
        background: linear-gradient(135deg, rgba(209, 53, 111, 0.05) 0%, rgba(212, 165, 116, 0.05) 100%);
        padding: 1.5rem;
        border-left: 4px solid var(--magenta);
        border-radius: 8px;
        margin-bottom: 1rem;
    }
    
    .section-title {
        color: var(--magenta);
        font-size: 1.5rem;
        border-bottom: 2px solid var(--gold);
        padding-bottom: 0.5rem;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# City coordinates for mapping
CITY_COORDINATES = {
    # North India
    "Delhi": (28.6139, 77.2090),
    "Agra": (27.1767, 78.0081),
    "Jaipur": (26.9124, 75.7873),
    "Amritsar": (31.6340, 74.8711),
    "Varanasi": (25.3268, 82.9856),
    "Shimla": (31.7725, 77.1740),
    "Manali": (32.2396, 77.1887),
    "Leh": (34.1526, 77.5771),
    "Rishikesh": (30.0887, 78.2672),
    "Haridwar": (29.9457, 78.1642),
    # West India
    "Mumbai": (19.0760, 72.8777),
    "Goa": (15.3008, 73.8207),
    "Udaipur": (24.5854, 73.7125),
    "Jodhpur": (26.2389, 73.0243),
    "Jaisalmer": (26.9124, 70.9185),
    "Ahmedabad": (23.0225, 72.5714),
    "Pune": (18.5204, 73.8567),
    # South India
    "Kochi": (9.9312, 76.2673),
    "Munnar": (10.5869, 77.0582),
    "Thiruvananthapuram": (8.5241, 76.9366),
    "Chennai": (13.0827, 80.2707),
    "Mysore": (12.2958, 76.6394),
    "Ooty": (11.4102, 76.6955),
    "Bengaluru": (12.9716, 77.5946),
    # East & Northeast
    "Kolkata": (22.5726, 88.3639),
    "Darjeeling": (27.0410, 88.2663),
    "Gangtok": (27.5330, 88.6109),
    "Kaziranga": (26.6307, 93.1872),
    # Nepal
    "Kathmandu": (27.7172, 85.3240),
    "Pokhara": (28.2096, 83.9856),
    "Chitwan": (27.5769, 84.4200),
    "Lumbini": (27.4926, 82.4997),
    # Bhutan
    "Thimphu": (27.5142, 89.6430),
    "Paro": (27.4283, 89.6233),
    "Punakha": (27.6039, 89.7535),
}

# Comprehensive itinerary database (scalable structure)
ITINERARIES_DB = {
    "India": {
        "Wildlife": {
            "Kaziranga Safari": {
                "cities": ["Guwahati", "Kaziranga"],
                "days_per_city": {"Guwahati": 1, "Kaziranga": 3},
                "description": "Rhino and tiger spotting in Assam's premier wildlife sanctuary",
                "themes": ["Wildlife", "Nature", "Adventure"],
                "best_season": "November to April"
            },
            "Tiger Quest Bandhavgarh": {
                "cities": ["Jabalpur", "Bandhavgarh"],
                "days_per_city": {"Jabalpur": 1, "Bandhavgarh": 4},
                "description": "Tiger spotting in Bandhavgarh National Park",
                "themes": ["Wildlife", "Adventure", "Photography"],
                "best_season": "November to June"
            },
            "Ranthambore Tigers": {
                "cities": ["Jaipur", "Sawai Madhopur"],
                "days_per_city": {"Jaipur": 2, "Sawai Madhopur": 3},
                "description": "Tiger safaris in Ranthambore National Park",
                "themes": ["Wildlife", "Adventure", "Photography"],
                "best_season": "October to June"
            },
            "Kanha Tigers": {
                "cities": ["Jabalpur", "Kanha"],
                "days_per_city": {"Jabalpur": 1, "Kanha": 3},
                "description": "Striped tigers and gaur in Kanha National Park",
                "themes": ["Wildlife", "Adventure"],
                "best_season": "November to May"
            },
            "Corbett National Park": {
                "cities": ["Delhi", "Nainital", "Corbett"],
                "days_per_city": {"Delhi": 1, "Nainital": 1, "Corbett": 3},
                "description": "Tigers and adventure in Uttarakhand",
                "themes": ["Wildlife", "Adventure", "Nature"],
                "best_season": "November to June"
            },
        },
        "Cultural": {
            "Golden Triangle": {
                "cities": ["Delhi", "Jaipur", "Agra"],
                "days_per_city": {"Delhi": 2, "Jaipur": 2, "Agra": 2},
                "description": "India's most iconic cultural route",
                "themes": ["Culture", "History", "Architecture"],
                "best_season": "October to March"
            },
            "Temple Trail Varanasi": {
                "cities": ["Delhi", "Varanasi"],
                "days_per_city": {"Delhi": 2, "Varanasi": 4},
                "description": "Sacred temples and spiritual experiences",
                "themes": ["Culture", "Spiritual", "History"],
                "best_season": "October to March"
            },
            "Rajasthan Royal": {
                "cities": ["Jaipur", "Jodhpur", "Udaipur"],
                "days_per_city": {"Jaipur": 3, "Jodhpur": 2, "Udaipur": 3},
                "description": "Royal palaces and forts of Rajasthan",
                "themes": ["Culture", "Architecture", "History"],
                "best_season": "October to March"
            },
            "South Indian Temples": {
                "cities": ["Chennai", "Madurai", "Thiruvananthapuram"],
                "days_per_city": {"Chennai": 2, "Madurai": 2, "Thiruvananthapuram": 2},
                "description": "Ancient Dravidian temples and traditions",
                "themes": ["Culture", "Spiritual", "History"],
                "best_season": "October to March"
            },
            "Kerala Backwater Culture": {
                "cities": ["Kochi", "Munnar", "Thiruvananthapuram"],
                "days_per_city": {"Kochi": 2, "Munnar": 2, "Thiruvananthapuram": 2},
                "description": "Spices, backwaters, and culture",
                "themes": ["Culture", "Nature", "Food"],
                "best_season": "September to May"
            },
        },
        "Beaches": {
            "Goa Paradise": {
                "cities": ["Mumbai", "Goa"],
                "days_per_city": {"Mumbai": 2, "Goa": 4},
                "description": "White sand beaches and Portuguese heritage",
                "themes": ["Beaches", "Water Sports", "Relaxation"],
                "best_season": "October to May"
            },
            "Kerala Beach Hop": {
                "cities": ["Kochi", "Munnar", "Alleppey", "Varkala"],
                "days_per_city": {"Kochi": 1, "Munnar": 2, "Alleppey": 2, "Varkala": 2},
                "description": "Beaches and backwaters of Kerala",
                "themes": ["Beaches", "Nature", "Relaxation"],
                "best_season": "August to March"
            },
            "Andaman Beaches": {
                "cities": ["Port Blair", "Havelock", "Neil Island"],
                "days_per_city": {"Port Blair": 1, "Havelock": 3, "Neil Island": 2},
                "description": "Tropical paradise and island adventures",
                "themes": ["Beaches", "Water Sports", "Nature"],
                "best_season": "November to May"
            },
            "Pondicherry Escape": {
                "cities": ["Chennai", "Pondicherry"],
                "days_per_city": {"Chennai": 1, "Pondicherry": 3},
                "description": "French colonial town and beach relaxation",
                "themes": ["Beaches", "Culture", "History"],
                "best_season": "October to March"
            },
        },
        "Adventure": {
            "Ladakh Extreme": {
                "cities": ["Delhi", "Leh", "Nubra Valley", "Pangong"],
                "days_per_city": {"Delhi": 1, "Leh": 2, "Nubra Valley": 2, "Pangong": 2},
                "description": "High altitude adventure and stunning landscapes",
                "themes": ["Adventure", "Trekking", "Nature"],
                "best_season": "June to September"
            },
            "Himalayan Trek": {
                "cities": ["Delhi", "Shimla", "Manali", "Rohtang"],
                "days_per_city": {"Delhi": 1, "Shimla": 2, "Manali": 3, "Rohtang": 1},
                "description": "Mountain trekking and alpine scenery",
                "themes": ["Adventure", "Trekking", "Nature"],
                "best_season": "April to October"
            },
            "Auli Skiing": {
                "cities": ["Delhi", "Auli"],
                "days_per_city": {"Delhi": 1, "Auli": 3},
                "description": "Skiing and snow sports in the Himalayas",
                "themes": ["Adventure", "Winter Sports"],
                "best_season": "December to March"
            },
            "Desert Safari": {
                "cities": ["Jaisalmer", "Sam Sand Dunes"],
                "days_per_city": {"Jaisalmer": 2, "Sam Sand Dunes": 2},
                "description": "Camel safaris and desert camps",
                "themes": ["Adventure", "Desert", "Nature"],
                "best_season": "October to March"
            },
        },
        "Spiritual": {
            "Yoga & Wellness": {
                "cities": ["Delhi", "Rishikesh"],
                "days_per_city": {"Delhi": 1, "Rishikesh": 5},
                "description": "Yoga, meditation, and spiritual retreats",
                "themes": ["Spiritual", "Wellness", "Yoga"],
                "best_season": "September to May"
            },
            "Sacred Pilgrimage": {
                "cities": ["Haridwar", "Varanasi", "Prayagraj"],
                "days_per_city": {"Haridwar": 2, "Varanasi": 3, "Prayagraj": 2},
                "description": "Holy Ganges and sacred sites",
                "themes": ["Spiritual", "Pilgrimage", "History"],
                "best_season": "October to March"
            },
        },
        "Festivals": {
            "Holi Celebration": {
                "cities": ["Delhi", "Jaipur", "Mathura"],
                "days_per_city": {"Delhi": 1, "Jaipur": 1, "Mathura": 2},
                "description": "Festival of colors celebration",
                "themes": ["Festival", "Culture", "Celebration"],
                "best_season": "March"
            },
            "Diwali": {
                "cities": ["Delhi", "Jaipur", "Varanasi"],
                "days_per_city": {"Delhi": 2, "Jaipur": 2, "Varanasi": 2},
                "description": "Festival of lights across India",
                "themes": ["Festival", "Culture", "Spiritual"],
                "best_season": "October-November"
            },
            "Pushkar Camel Fair": {
                "cities": ["Jaipur", "Pushkar"],
                "days_per_city": {"Jaipur": 2, "Pushkar": 3},
                "description": "Colorful camel fair and cultural celebration",
                "themes": ["Festival", "Culture", "Markets"],
                "best_season": "November-December"
            },
        },
    },
    "Nepal": {
        "Adventure": {
            "Everest Base Camp Trek": {
                "cities": ["Kathmandu", "Lukla", "EBC"],
                "days_per_city": {"Kathmandu": 2, "Lukla": 1, "EBC": 10},
                "description": "Trek to the base of Mount Everest",
                "themes": ["Adventure", "Trekking", "Mountain"],
                "best_season": "September to November, March to May"
            },
            "Annapurna Circuit": {
                "cities": ["Kathmandu", "Pokhara", "Annapurna"],
                "days_per_city": {"Kathmandu": 2, "Pokhara": 2, "Annapurna": 12},
                "description": "World's most popular trekking circuit",
                "themes": ["Adventure", "Trekking", "Mountain"],
                "best_season": "September to November, March to May"
            },
            "Kathmandu Valley Trek": {
                "cities": ["Kathmandu", "Nagarkot", "Bhaktapur"],
                "days_per_city": {"Kathmandu": 2, "Nagarkot": 2, "Bhaktapur": 2},
                "description": "Cultural and nature trekking near Kathmandu",
                "themes": ["Adventure", "Culture", "Hiking"],
                "best_season": "Year-round"
            },
        },
        "Cultural": {
            "Kathmandu Valley Temples": {
                "cities": ["Kathmandu", "Patan", "Bhaktapur"],
                "days_per_city": {"Kathmandu": 3, "Patan": 2, "Bhaktapur": 2},
                "description": "UNESCO temples and Newar architecture",
                "themes": ["Culture", "History", "Architecture"],
                "best_season": "Year-round"
            },
            "Pokhara Exploration": {
                "cities": ["Kathmandu", "Pokhara"],
                "days_per_city": {"Kathmandu": 2, "Pokhara": 3},
                "description": "Lake city and mountain views",
                "themes": ["Culture", "Nature", "Relaxation"],
                "best_season": "September to November, March to May"
            },
        },
        "Spiritual": {
            "Pilgrimage Journey": {
                "cities": ["Kathmandu", "Lumbini", "Chitwan"],
                "days_per_city": {"Kathmandu": 2, "Lumbini": 2, "Chitwan": 2},
                "description": "Buddhist pilgrimage sites",
                "themes": ["Spiritual", "Pilgrimage", "Culture"],
                "best_season": "Year-round"
            },
        },
    },
    "Bhutan": {
        "Cultural": {
            "Thimphu to Paro": {
                "cities": ["Thimphu", "Paro"],
                "days_per_city": {"Thimphu": 3, "Paro": 3},
                "description": "Capital and sacred valley exploration",
                "themes": ["Culture", "Spiritual", "Architecture"],
                "best_season": "March to May, September to November"
            },
            "Punakha Valley": {
                "cities": ["Thimphu", "Punakha", "Paro"],
                "days_per_city": {"Thimphu": 2, "Punakha": 3, "Paro": 2},
                "description": "Dzongs and traditional villages",
                "themes": ["Culture", "History", "Architecture"],
                "best_season": "March to May, September to November"
            },
        },
        "Spiritual": {
            "Tiger's Nest & Monasteries": {
                "cities": ["Paro", "Punakha", "Bumthang"],
                "days_per_city": {"Paro": 3, "Punakha": 2, "Bumthang": 2},
                "description": "Sacred Buddhist sites and monasteries",
                "themes": ["Spiritual", "Culture", "Trekking"],
                "best_season": "March to May, September to November"
            },
            "Spiritual Retreat": {
                "cities": ["Thimphu", "Paro", "Punakha"],
                "days_per_city": {"Thimphu": 2, "Paro": 3, "Punakha": 2},
                "description": "Meditation and spiritual practices",
                "themes": ["Spiritual", "Wellness", "Culture"],
                "best_season": "March to May, September to November"
            },
        },
        "Adventure": {
            "Druk Path Trek": {
                "cities": ["Thimphu", "Paro"],
                "days_per_city": {"Thimphu": 2, "Paro": 5},
                "description": "High altitude mountain trekking",
                "themes": ["Adventure", "Trekking", "Nature"],
                "best_season": "March to May, September to October"
            },
        },
    }
}

# Initialize session state
if 'itinerary' not in st.session_state:
    st.session_state.itinerary = None
if 'generated' not in st.session_state:
    st.session_state.generated = False
if 'custom_days' not in st.session_state:
    st.session_state.custom_days = {}

# Header with Logo
col1, col2, col3 = st.columns([1, 2, 1])
with col1:
    st.image("https://raw.githubusercontent.com/indiareisen/india-reisen-itinerary/main/logo.png", width=80)
with col2:
    st.markdown('<div class="main-header">India Reisen</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Full Service Immersive Journeys</div>', unsafe_allow_html=True)
with col3:
    st.image("https://raw.githubusercontent.com/indiareisen/india-reisen-itinerary/main/logo.png", width=80)

st.divider()

# Selection section
st.markdown('<div class="section-title">Plan Your Journey</div>', unsafe_allow_html=True)

col1, col2, col3 = st.columns(3)

with col1:
    country = st.selectbox("Select Country", list(ITINERARIES_DB.keys()))

with col2:
    if country in ITINERARIES_DB:
        themes = list(ITINERARIES_DB[country].keys())
        selected_theme = st.selectbox("Select Theme", themes)

with col3:
    if country in ITINERARIES_DB and selected_theme in ITINERARIES_DB[country]:
        itineraries = list(ITINERARIES_DB[country][selected_theme].keys())
        selected_itinerary = st.selectbox("Select Itinerary", itineraries)

st.divider()

# Get selected itinerary details
if country in ITINERARIES_DB and selected_theme in ITINERARIES_DB[country] and selected_itinerary in ITINERARIES_DB[country][selected_theme]:
    itin = ITINERARIES_DB[country][selected_theme][selected_itinerary]
    
    st.markdown('<div class="section-title">Customize Days per City</div>', unsafe_allow_html=True)
    
    # Display and customize days for each city
    cols = st.columns(len(itin["cities"]))
    custom_days = {}
    
    for idx, city in enumerate(itin["cities"]):
        with cols[idx]:
            default_days = itin["days_per_city"].get(city, 2)
            custom_days[city] = st.number_input(
                f"Days in {city}",
                min_value=1,
                max_value=30,
                value=default_days,
                key=f"days_{city}"
            )
    
    st.session_state.custom_days = custom_days
    total_days = sum(custom_days.values())
    st.info(f"📅 **Total Journey Duration: {total_days} days**")
    
    # Display itinerary info
    st.markdown('<div class="section-title">Journey Overview</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        st.write(f"**Description:** {itin.get('description', 'Custom journey')}")
    with col2:
        st.write(f"**Best Season:** {itin.get('best_season', 'Year-round')}")
    
    # Draw routing map
    st.markdown('<div class="section-title">Journey Route Map</div>', unsafe_allow_html=True)
    
    # Create map centered on first city
    first_city = itin["cities"][0]
    if first_city in CITY_COORDINATES:
        first_coords = CITY_COORDINATES[first_city]
        
        m = folium.Map(
            location=first_coords,
            zoom_start=6,
            tiles="OpenStreetMap"
        )
        
        # Add markers for each city
        colors = ['red', 'blue', 'green', 'purple', 'orange', 'darkred', 'darkblue', 'darkgreen']
        
        for idx, city in enumerate(itin["cities"]):
            if city in CITY_COORDINATES:
                coords = CITY_COORDINATES[city]
                color = colors[idx % len(colors)]
                days = custom_days.get(city, itin["days_per_city"].get(city, 1))
                
                folium.Marker(
                    location=coords,
                    popup=f"<b>{city}</b><br>Days: {days}",
                    tooltip=city,
                    icon=folium.Icon(color=color, icon='info-sign', prefix='glyphicon')
                ).add_to(m)
        
        # Draw route lines
        route_coords = []
        for city in itin["cities"]:
            if city in CITY_COORDINATES:
                route_coords.append(CITY_COORDINATES[city])
        
        if len(route_coords) > 1:
            folium.PolyLine(
                route_coords,
                color='red',
                weight=2,
                opacity=0.7,
                popup='Journey Route'
            ).add_to(m)
        
        st_folium(m, width=1200, height=400)
    
    st.divider()
    
    # Generate button
    if st.button("✨ Generate Full Itinerary", use_container_width=True, type="primary"):
        st.session_state.itinerary = itin
        st.session_state.generated = True
        st.session_state.selected_country = country
        st.session_state.selected_theme = selected_theme
        st.session_state.selected_itinerary = selected_itinerary
        st.success("✅ Itinerary generated successfully!")

# Display generated itinerary
if st.session_state.generated and st.session_state.itinerary:
    itin = st.session_state.itinerary
    
    st.markdown('<div class="section-title">Your Customized Journey</div>', unsafe_allow_html=True)
    
    st.markdown(f"### {st.session_state.selected_itinerary}")
    st.markdown(f"**Country:** {st.session_state.selected_country} | **Theme:** {st.session_state.selected_theme}")
    st.markdown(f"**Route:** {' → '.join(itin['cities'])}")
    st.markdown(f"**Total Duration:** {sum(st.session_state.custom_days.values())} days")
    
    # Day-by-day summary
    st.markdown('<div class="section-title">Journey Breakdown</div>', unsafe_allow_html=True)
    
    day_counter = 1
    for city in itin["cities"]:
        days = st.session_state.custom_days.get(city, itin["days_per_city"].get(city, 1))
        
        with st.expander(f"**{city}** - Days {day_counter} to {day_counter + days - 1} ({days} days)"):
            st.write(f"📍 **City:** {city}")
            st.write(f"🗓️ **Duration:** {days} days")
            st.write(f"🌍 **Theme:** {', '.join(itin.get('themes', ['Travel']))}")
            st.write(f"🎯 **Best Season:** {itin.get('best_season', 'Year-round')}")
            st.write("""
            #### Suggested Activities:
            - Local exploration and cultural immersion
            - Guided tours of major attractions
            - Local cuisine and dining experiences
            - Shopping at markets and bazaars
            - Evening entertainment and relaxation
            """)
        
        day_counter += days
    
    # Export
    st.divider()
    st.markdown('<div class="section-title">Export Your Itinerary</div>', unsafe_allow_html=True)
    
    def generate_docx(itin, country, theme, title):
        doc = Document()
        
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("🌸 INDIA REISEN")
        title_run.font.size = Pt(28)
        title_run.font.color.rgb = RGBColor(209, 53, 111)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(title).style = 'Heading 1'
        doc.add_paragraph(f"Country: {country}").style = 'Heading 2'
        doc.add_paragraph(f"Theme: {theme}")
        doc.add_paragraph(f"Total Duration: {sum(st.session_state.custom_days.values())} days")
        
        doc.add_heading('Journey Route', level=2).runs[0].font.color.rgb = RGBColor(209, 53, 111)
        doc.add_paragraph(' → '.join(itin['cities']))
        
        doc.add_heading('Daily Breakdown', level=2).runs[0].font.color.rgb = RGBColor(209, 53, 111)
        
        day_counter = 1
        for city in itin["cities"]:
            days = st.session_state.custom_days.get(city, itin["days_per_city"].get(city, 1))
            doc.add_heading(f"{city} - {days} days", level=3)
            doc.add_paragraph(f"Days: {day_counter} to {day_counter + days - 1}")
            doc.add_paragraph("Includes local exploration, guided tours, cultural immersion, and dining experiences")
            day_counter += days
        
        doc.add_paragraph()
        footer = doc.add_paragraph("🌸 Book your immersive journey with India Reisen today!")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    if st.button("📝 Download as DOCX", use_container_width=True):
        docx_buffer = generate_docx(itin, st.session_state.selected_country, st.session_state.selected_theme, st.session_state.selected_itinerary)
        st.download_button(
            label="📥 Download DOCX",
            data=docx_buffer,
            file_name=f"India_Reisen_{st.session_state.selected_itinerary}_{datetime.now().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

st.divider()
st.markdown("""
---
**India Reisen** | Full Service Immersive Journeys  
🌐 Destinations: India • Nepal • Bhutan • Tibet • Sri Lanka  
📱 Follow us: Instagram • Facebook • LinkedIn • YouTube

*Created with ✨ for travelers seeking authentic cultural immersion*
""")
